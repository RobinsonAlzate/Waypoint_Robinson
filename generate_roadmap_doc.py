import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def create_document():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Waypoint Project: Executive Deliverable Roadmap & Technical Guide")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Prepared for Professor Gabriel Albu | Application Programming Course (Summer 2026)")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section Helper
    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 102, 153)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.name = "Arial"
        r2 = p.add_run(text)
        r2.font.name = "Arial"
        p.paragraph_format.space_after = Pt(2)
        return p

    # Executive Summary
    add_heading_1("1. Executive Summary")
    p = doc.add_paragraph(
        "This document serves as the master instructional manual and technical roadmap for the Waypoint trail-finder and trip-planner application. The project spans an 8-week developmental lifecycle (Weeks 7 through 14), transitioning from a pure-Python object-oriented domain engine to a fully operational Django web application.")
    p.runs[0].font.name = "Arial"

    # Directory Structure
    add_heading_1("2. Application Directory Architecture")
    p = doc.add_paragraph(
        "The application resides under Documents\\Cloud Computing\\Applications Programming\\Lecture1\\pythonProject\\pythonProject\\waypoint:")
    p.runs[0].font.name = "Arial"

    add_bullet("env/",
               "Isolated virtual environment housing dependencies (Django, docx, lxml, sqlparse, tzdata) and execution scripts.")
    add_bullet("static/css/style.css/", "Global cascading style sheet managing custom layout and design aesthetics.")
    add_bullet("templates/",
               "Master layout (base.html), pages (home.html, report_form.html, search.html, thank_you.html), partial components (navbar.html, footer.html), and trail module views (catalog.html, park_detail.html, trail_detail.html).")

    # Phase 1 & 2 Roadmap
    add_heading_1("3. Week-by-Week Deliverables & Testing Guide")

    phases = [
        ("Week 7 — Domain Model: Classes & Objects (10%)",
         "waypoint/waypoint_core/trail.py, distance.py, test_domain.py",
         "python waypoint/waypoint_core/test_domain.py",
         "Validates that attributes are properly encapsulated within class definitions and instance methods function correctly without external dependencies."),

        ("Week 8 — Domain Model: Inheritance, Polymorphism & Operators (12%)",
         "waypoint/waypoint_core/trail.py, itinerary.py, test_domain.py",
         "python waypoint/waypoint_core/test_domain.py",
         "Assesses advanced object-oriented programming principles, confirming polymorphism and operator overloading behave correctly across trail types."),

        ("Week 9 — Django Project Setup (6%)",
         "waypoint/waypoint/settings.py, urls.py, manage.py",
         "python waypoint/manage.py runserver",
         "Confirms environment and Django core are correctly initialized. Note: Ensure local testing requests use HTTP explicitly, as modern web browsers may attempt to force HTTPS connections."),

        ("Week 10 — Views, URLs & The Report Form (12%)",
         "waypoint/waypoint/views.py, urls.py, report_form.html, thank_you.html",
         "python waypoint/manage.py runserver (navigate to /report/)",
         "Verifies correct request/response handling, proper URL resolution, and inclusion of CSRF tokens on form submissions."),

        ("Week 11 — The Catalog Templates & Static Assets (12%)",
         "waypoint/templates/base.html, partials/, trails/catalog.html, static/css/style.css",
         "Navigate to /catalog/ in browser",
         "Assesses front-end organization, DRY template inheritance principles, and proper static file loading."),

        ("Week 12 — ORM, Models & Admin Integration (14%)",
         "waypoint/trails/models.py, admin.py, migrations/0001_initial.py",
         "python waypoint/manage.py runserver (navigate to /admin/)",
         "Validates database schema design, migration history integrity, and admin panel registration."),

        ("Week 13 — Relational Modeling: ForeignKey Links (12%)",
         "waypoint/trails/models.py, management/commands/seed_trails.py",
         "python waypoint/manage.py seed_trails",
         "Ensures proper database normalization, referential integrity, and automated data population capability."),

        ("Week 14 — Hardening, Documentation & Handoff (12%)",
         "README.md, TESTING_GUIDE.md, generate_roadmap_doc.py",
         "python waypoint/generate_roadmap_doc.py",
         "Validates end-to-end execution of documentation scripts and proves project readiness for final submission.")
    ]

    for title_text, files, test_cmd, reason in phases:
        add_heading_2(title_text)
        add_bullet("Files to Show: ", files)
        add_bullet("How to Test: ", test_cmd)
        add_bullet("Grading Rationale: ", reason)

    doc.save("Waypoint_Executive_Deliverable_Roadmap2.docx")
    print("Successfully generated Waypoint_Executive_Deliverable_Roadmap2.docx")


if __name__ == "__main__":
    create_document()